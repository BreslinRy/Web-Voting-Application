from docx import Document
import csv
from docx2pdf import convert


def genReceipt(VoterData):
    #init voter ID and Document
    voterID = 0
    receipt = Document()
    #setup header of Receipt, bolded and center aligned
    header = receipt.add_heading('Voters Receipt')
    header.bold = True
    header.alignment = 1
    #open CSV file to Read
    with open(VoterData) as csvFile:
        #init line counter and Csv reader
        lineCount = 1
        csvReader = csv.reader(csvFile, delimiter=',')
        #loop through each element of the csv
        for row in csvReader:
            #if first row of data, set up opening paragraph grabbing the Voter ID number and adding it the the document
            if lineCount == 1:
                voterID = row[0]
                intro = receipt.add_paragraph('This receipt corresponds to voter ID: ')
                intro.add_run(voterID)
                intro.add_run('.')
            #determine if ballot line is out of order
            if lineCount == int(row[1]):
                #determine if line is a proposition result and if true add prop result
                if row[2] == 'yes' or row[2] == 'no':
                    vote = receipt.add_paragraph('You have voted ')
                    vote.add_run(row[2])
                    vote.add_run(' for proposition ')
                    vote.add_run(row[3])
                    vote.add_run('.')
                #if not a prop add to doc as a canidate result
                else:
                    vote = receipt.add_paragraph('You have voted for ')
                    vote.add_run(row[2])
                    vote.add_run(' for the position of ')
                    vote.add_run(row[3])
                    vote.add_run('.')
                #increment line count
                lineCount = lineCount + 1
            #if ballot line is out of order rewrite doc as an invaild ballot error notification
            else:
                receipt = Document()
                header = receipt.add_heading('INVALID BALLOT')
                header.bold = True
                header.alignment = 1
                error = receipt.add_paragraph('An error has been encountered with your ballot please contact your polling station to resolve the issue for Voter ID: ')
                error.add_run(voterID)
                error.add_run('.')
                receipt.save('VotingReceipt.docx')
                convert("VotingReceipt.docx")
                #return failed
                return "Failed due to inconsistent ballot data"
    #add final paragraph, then save and convert file to a pdf.
    receipt.add_paragraph('Thank you for voting in this years election. Your votes have been recorded and this document acts as proof. Please save and print this file to maintain a paper trail of this interaction.')
    receipt.save('VotingReceipt.docx')
    convert("VotingReceipt.docx")
    # return success
    return "Success"